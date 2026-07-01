"""Generate publication-ready Word tables for all four outcomes.

Outputs results/tables.docx with:
  Table 1   Bivariate correlations (primary outcome: all adult brain tumours).
  Table 2   Bivariate correlations (secondary: glioblastoma).
  Table 3   Bivariate correlations (secondary: diffuse and anaplastic astrocytoma).
  Table 4   Bivariate correlations (secondary: oligodendroglioma).
  Table 5   Boxplot summary statistics for binary indicators, all four outcomes.
  Table 6   Multivariable OLS regression (primary outcome).
  Table S1  Country-level analytic dataset excerpt with provenance.
  Table S2  Sensitivity: high-quality CONCORD subset, all four outcomes.
"""
from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


OUTCOMES = [
    ("brain_all",     "all adult brain tumours"),
    ("glioblastoma",  "glioblastoma"),
    ("diffuse_anap",  "diffuse and anaplastic astrocytoma"),
    ("oligo",         "oligodendroglioma"),
]


def fmt(v, digits=3):
    if pd.isna(v):
        return "."
    if isinstance(v, (int,)):
        return f"{v:d}"
    return f"{v:.{digits}f}"


def add_table(doc: Document, df: pd.DataFrame, title: str, footer: str = "") -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, col in enumerate(df.columns):
            cell = table.rows[i].cells[j]
            cell.text = str(row[col])
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if footer:
        p = doc.add_paragraph()
        run = p.add_run(footer)
        run.italic = True
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def correlation_df(corr_csv: Path) -> pd.DataFrame:
    df = pd.read_csv(corr_csv)
    rows = []
    for _, r in df.iterrows():
        rows.append({
            "Indicator": r["indicator"],
            "n": int(r["n"]),
            "Pearson r": fmt(r["pearson_r"]),
            "95% CI (boot)": f"{fmt(r['pearson_lo'])} to {fmt(r['pearson_hi'])}",
            "p (raw)": fmt(r["pearson_p"]),
            "q (BH)": fmt(r["pearson_q_bh_within_outcome"]),
            "Spearman rho": fmt(r["spearman_r"]),
            "n for 80% power": fmt(r["n_required_for_80pct_power"], 0)
                if pd.notna(r["n_required_for_80pct_power"]) else ".",
        })
    return pd.DataFrame(rows)


def binary_combined_df(results_dir: Path) -> pd.DataFrame:
    rows = []
    for outcome_id, label in OUTCOMES:
        df = pd.read_csv(results_dir / f"binary_indicator_summary__{outcome_id}.csv")
        for _, r in df.iterrows():
            rows.append({
                "Outcome": label,
                "Binary indicator": r["indicator"],
                "n=0 (median, IQR)": (f"{int(r['n_0'])} ({fmt(r['median_0'], 1)}, "
                                      f"{fmt(r['q1_0'], 1)}-{fmt(r['q3_0'], 1)})"),
                "n=1 (median, IQR)": (f"{int(r['n_1'])} ({fmt(r['median_1'], 1)}, "
                                      f"{fmt(r['q1_1'], 1)}-{fmt(r['q3_1'], 1)})"),
                "p (Mann-Whitney)": fmt(r["mann_whitney_p"]),
            })
    return pd.DataFrame(rows)


def regression_df(reg_csv: Path) -> pd.DataFrame:
    df = pd.read_csv(reg_csv)
    rows = []
    for _, r in df.iterrows():
        rows.append({
            "Indicator (continuous)": r["indicator"],
            "n": int(r["n"]),
            "Beta (HC3 SE)": f"{fmt(r['beta'])} ({fmt(r['se'])})",
            "p (indicator)": fmt(r["p"]),
            "Beta log GDP pc": fmt(r["beta_covariate"]),
            "p (log GDP pc)": fmt(r["p_covariate"]),
            "Model R^2": fmt(r["r2"]),
        })
    return pd.DataFrame(rows)


def sensitivity_df(results_dir: Path) -> pd.DataFrame:
    rows = []
    for outcome_id, label in OUTCOMES:
        df = pd.read_csv(results_dir / f"sensitivity_high_quality__{outcome_id}.csv")
        for _, r in df.iterrows():
            rows.append({
                "Outcome": label,
                "Indicator": r["indicator"],
                "n": int(r["n"]),
                "Pearson r": fmt(r["pearson_r"]),
                "95% CI (boot)": f"{fmt(r['pearson_lo'])} to {fmt(r['pearson_hi'])}",
                "Spearman rho": fmt(r["spearman_r"]),
            })
    return pd.DataFrame(rows)


def country_df(analytic_csv: Path) -> pd.DataFrame:
    df = pd.read_csv(analytic_csv)
    cols = [
        ("iso3", "ISO3"),
        ("country", "Country"),
        ("brain_adult_5yr_pct", "Brain (all)"),
        ("glioblastoma_5yr_pct", "GBM"),
        ("oligodendroglioma_5yr", "Oligo"),
        ("oop_pct_che", "OOP%CHE"),
        ("nurses_per_1000", "Nurses/1k"),
        ("radiotherapy_units_per_million", "RT/million"),
        ("oral_morphine_available_2013", "Morph 2013"),
        ("right_to_health_in_constitution", "RtH const."),
        ("icescr_ratified", "ICESCR"),
    ]
    df_s = df[[c for c, _ in cols]].copy()
    df_s.columns = [n for _, n in cols]
    for col in df_s.columns:
        df_s[col] = df_s[col].apply(
            lambda v: fmt(v, 2) if isinstance(v, float) else (str(v) if pd.notna(v) else ".")
        )
    return df_s


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--results_dir", required=True, type=Path)
    ap.add_argument("--analytic", required=True, type=Path)
    ap.add_argument("--out_docx", required=True, type=Path)
    args = ap.parse_args()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("Right to health and adult brain tumour survival: tables")
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph(
        "Tables generated from the analytic dataset described in the README. Survival "
        "data: CONCORD-3 main analysis (Allemani et al., Lancet 2018) and CONCORD-3 "
        "brain histology analysis (Girardi et al., Neuro-Oncology 2023). Indicator "
        "data: World Bank Open Data API, WHO Global Health Observatory, Backman et al. "
        "Lancet 2008, UN Treaty Collection. Period of accrual for survival: patients "
        "diagnosed 2010-14, age 15-99 at diagnosis."
    ).runs[0].font.size = Pt(9)

    titles = {
        "brain_all":    "Table 1. Bivariate correlations: all adult brain tumours, 5-year net survival 2010-14.",
        "glioblastoma": "Table 2. Bivariate correlations: glioblastoma, 5-year net survival 2010-14.",
        "diffuse_anap": "Table 3. Bivariate correlations: diffuse and anaplastic astrocytoma, 5-year net survival 2010-14.",
        "oligo":        "Table 4. Bivariate correlations: oligodendroglioma, 5-year net survival 2010-14.",
    }
    footer = ("Bootstrap 95% CIs from 2000 resamples (seed 20260426). q-values from "
              "Benjamini-Hochberg FDR over the 9 continuous indicator tests. Power "
              "calculated by Fisher-z (alpha=0.05).")
    for outcome_id, _label in OUTCOMES:
        add_table(
            doc,
            correlation_df(args.results_dir / f"correlation_results__{outcome_id}.csv"),
            titles[outcome_id], footer=footer,
        )

    add_table(
        doc, binary_combined_df(args.results_dir),
        "Table 5. Distribution of survival across binary right-to-health indicators "
        "(four outcomes).",
        footer="IQR = inter-quartile range. Mann-Whitney U test, two-sided.",
    )

    add_table(
        doc, regression_df(args.results_dir / "regression_results__brain_all.csv"),
        "Table 6. Multivariable OLS for the primary outcome (all adult brain tumours): "
        "5-year survival regressed on each indicator with log-GDP per capita PPP as a "
        "covariate.",
        footer="HC3 robust standard errors. Effective n varies by indicator availability.",
    )

    add_table(
        doc, sensitivity_df(args.results_dir),
        "Table S1. Sensitivity analysis: bivariate correlations restricted to the "
        "high-quality CONCORD-3 subset (100% national coverage, no flags).",
    )

    add_table(
        doc, country_df(args.analytic),
        "Table S2. Country-level analytic dataset (excerpt of key columns).",
    )

    doc.save(args.out_docx)
    print(f"Wrote {args.out_docx}")


if __name__ == "__main__":
    main()
